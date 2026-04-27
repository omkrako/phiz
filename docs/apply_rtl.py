"""
Post-processes a DOCX file to apply RTL direction throughout:
  - Document default paragraph/run properties
  - All named styles (Normal, Heading 1-6, List Paragraph, etc.)
  - Every paragraph and run in the body and tables
"""
import sys
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def _set_bidi_on_pPr(pPr):
    """Add <w:bidi/> and <w:jc val="right"/> to a paragraph properties element."""
    if pPr.find(qn('w:bidi')) is None:
        bidi = OxmlElement('w:bidi')
        bidi.set(qn('w:val'), '1')
        pPr.insert(0, bidi)
    jc = pPr.find(qn('w:jc'))
    if jc is None:
        jc = OxmlElement('w:jc')
        pPr.append(jc)
    jc.set(qn('w:val'), 'right')


def _set_rtl_on_rPr(rPr):
    """Add <w:rtl/> to a run properties element."""
    if rPr.find(qn('w:rtl')) is None:
        rtl = OxmlElement('w:rtl')
        rtl.set(qn('w:val'), '1')
        rPr.append(rtl)


def _apply_to_para(para):
    _set_bidi_on_pPr(para._p.get_or_add_pPr())
    for run in para.runs:
        _set_rtl_on_rPr(run._r.get_or_add_rPr())


def set_rtl(doc_path, out_path):
    doc = Document(doc_path)

    # 1. Document default paragraph properties (w:docDefaults > w:pPrDefault)
    docDefaults = doc.styles.element.find(qn('w:docDefaults'))
    if docDefaults is not None:
        pPrDefault = docDefaults.find(qn('w:pPrDefault'))
        if pPrDefault is None:
            pPrDefault = OxmlElement('w:pPrDefault')
            docDefaults.insert(0, pPrDefault)
        pPr = pPrDefault.find(qn('w:pPr'))
        if pPr is None:
            pPr = OxmlElement('w:pPr')
            pPrDefault.append(pPr)
        _set_bidi_on_pPr(pPr)

        rPrDefault = docDefaults.find(qn('w:rPrDefault'))
        if rPrDefault is None:
            rPrDefault = OxmlElement('w:rPrDefault')
            docDefaults.append(rPrDefault)
        rPr = rPrDefault.find(qn('w:rPr'))
        if rPr is None:
            rPr = OxmlElement('w:rPr')
            rPrDefault.append(rPr)
        _set_rtl_on_rPr(rPr)

    # 2. All named styles
    for style in doc.styles:
        try:
            el = style.element
            # Paragraph styles
            pPr = el.find(qn('w:pPr'))
            if pPr is not None:
                _set_bidi_on_pPr(pPr)
            # Run properties within styles
            rPr = el.find(qn('w:rPr'))
            if rPr is not None:
                _set_rtl_on_rPr(rPr)
        except Exception:
            pass

    # 3. Every paragraph in the body
    for para in doc.paragraphs:
        _apply_to_para(para)

    # 4. Every paragraph in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _apply_to_para(para)

    doc.save(out_path)
    print(f"Saved RTL document: {out_path}")


if __name__ == "__main__":
    if len(sys.argv) < 3:
        print("Usage: apply_rtl.py input.docx output.docx")
        sys.exit(1)
    set_rtl(sys.argv[1], sys.argv[2])
